import json
import docx
from simplify_docx import simplify
import os


def mark_titles(doc):
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = para.style.name.split(' ')[-1]  # 获取标题级别
            para.add_run(f' [Title{level}]')  # 标记


def parse_table(table_data):
    columns = []
    data = []

    for i, row in enumerate(table_data):
        if row['TYPE'] == 'table-row':
            row_data = []
            for cell in row['VALUE']:
                cell_texts = []
                for paragraph in cell['VALUE']:
                    if paragraph['TYPE'] == 'paragraph':
                        for text in paragraph['VALUE']:
                            if text['TYPE'] == 'text':
                                cell_texts.append(text['VALUE'])
                cell_content = ' '.join(cell_texts).strip()
                row_data.append(cell_content)

            if i == 0:
                # 第一行作为表头
                columns = row_data
            else:
                # 后续行为表格数据
                data.append(row_data)

    return {"columns": columns, "data": data}


def process_json_structure(data):
    result = {}
    stack = [result]
    previous_level = 0  # 跟踪上一个标题的层级

    for item in data['VALUE'][0]['VALUE']:
        if item['TYPE'] == 'paragraph':
            text = item['VALUE'][0]['VALUE']
            if '[Title' in text:
                # 提取标题级别和标题内容
                level = int(text.split('[Title')[-1][0])
                title = text.split('[Title')[0].strip()

                # 如果当前标题是n+2级，而没有n+1级标题，自动补充n+1级为null
                if level > previous_level + 1:
                    for l in range(previous_level + 1, level):
                        placeholder = {"content": None}
                        stack[-1][f"Default_Level_{l}"] = placeholder
                        stack.append(placeholder)

                # 确保堆栈保持正确的层级
                while len(stack) > level:
                    stack.pop()

                current_section = {"paragraphs": []}
                stack[-1][title] = current_section
                stack.append(current_section)
                previous_level = level

            else:
                # 添加当前自然段落到当前标题下
                if 'paragraphs' in stack[-1]:
                    stack[-1]['paragraphs'].append(text.strip())
                else:
                    stack[-1]['paragraphs'] = [text.strip()]

        elif item['TYPE'] == 'table':
            table = parse_table(item['VALUE'])
            if 'tables' not in stack[-1]:
                stack[-1]['tables'] = []
            stack[-1]['tables'].append(table)

    return result



def main():
    # 输入和输出路径
    input_docx_path = "../files/doctest.docx"
    output_json_path = os.path.join("../output", f"{os.path.splitext(os.path.basename(input_docx_path))[0]}-docx.json")

    # 读取 docx
    my_doc = docx.Document(input_docx_path)

    # 预处理-标记标题
    mark_titles(my_doc)

    # 将 docx 转换为初步 JSON
    my_doc_as_json = simplify(my_doc)

    # 二次处理-将 JSON 数据结构化为多级标题并保留自然段落和表格
    processed_json = process_json_structure(my_doc_as_json)

    with open(output_json_path, "w", encoding="utf-8") as final_json_file:
        json.dump(processed_json, final_json_file, ensure_ascii=False, indent=4)


if __name__ == "__main__":
    main()
